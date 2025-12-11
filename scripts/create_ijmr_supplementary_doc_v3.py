#!/usr/bin/env python3
"""
Assemble a comprehensive IJMR supplementary DOCX containing extended tables,
state-level outputs, and supporting figures. This script leaves prior artifacts
untouched and writes to the v3 submission directory for easy upload.
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

PROCESSED_DIR = Path("data/processed")
OUTPUT_DIR = Path("IJMR_Submission_v3_best")
OUTPUT_PATH = OUTPUT_DIR / "ijmr_supplementary_material.docx"

FIGURES = [
    (
        "Supplementary Figure S1. Bayesian posterior for patient delay.",
        Path("submission_package/figures/supplementary_bayesian_forest_patient_delay_(days).png"),
    ),
    (
        "Supplementary Figure S2. Bayesian posterior for diagnostic delay.",
        Path("submission_package/figures/supplementary_bayesian_forest_diagnostic_delay_(days).png"),
    ),
    (
        "Supplementary Figure S3. PCA loadings heatmap.",
        Path("submission_package/figures/supplementary_pca_loadings_heatmap_delay_determinants.png"),
    ),
    (
        "Supplementary Figure S4. PCA biplot of delay determinants.",
        Path("submission_package/figures/supplementary_pca_biplot_delay_determinants.png"),
    ),
]


def set_layout(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def add_table(doc: Document, caption: str, columns: Sequence[str], rows: Iterable[Sequence[str]]) -> None:
    para = doc.add_paragraph(caption)
    para.runs[0].bold = True
    table = doc.add_table(rows=len(rows) + 1, cols=len(columns))
    table.style = "Light List Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, header in enumerate(columns):
        cell = table.cell(0, j)
        cell.text = str(header)
        cell.paragraphs[0].runs[0].bold = True
    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            table.cell(i, j).text = str(value)
    doc.add_paragraph()


def add_component_scores(doc: Document, df: pd.DataFrame) -> None:
    columns = ["State", "PC1", "PC2", "PC3"]
    rows = [
        [
            row["state"].title(),
            f"{row['PC1']:.2f}",
            f"{row['PC2']:.2f}",
            f"{row['PC3']:.2f}",
        ]
        for _, row in df.sort_values("state").iterrows()
    ]
    add_table(doc, "Supplementary Table S1. PCA component scores for Indian states.", columns, rows)


def add_loadings(doc: Document, df: pd.DataFrame) -> None:
    df = df.rename(columns={"Unnamed: 0": "determinant"})
    components = [c for c in df.columns if c.startswith("PC")]
    columns = ["Determinant"] + components[:3]
    rows = [
        [row["determinant"]] + [f"{row[comp]:.3f}" for comp in components[:3]]
        for _, row in df.iterrows()
    ]
    add_table(doc, "Supplementary Table S2. PCA loadings for key determinants.", columns, rows)


def add_dag_paths(doc: Document, df: pd.DataFrame) -> None:
    columns = ["Outcome", "Source", "Path", "Evidence strength", "Avg score"]
    rows = [
        [
            row["outcome"],
            row["source"],
            row["path"],
            row["evidence_strengths"],
            f"{row['avg_evidence_score']:.2f}",
        ]
        for _, row in df.iterrows()
    ]
    add_table(doc, "Supplementary Table S3. Enumerated DAG pathways for detection delay.", columns, rows)


def add_state_rankings(doc: Document, df: pd.DataFrame) -> None:
    columns = ["Rank", "State", "Composite score", "PN ratio", "IN ratio", "Symptomatic no-care (%)", "Poverty (%)"]
    rows = [
        [
            int(row["priority_rank"]),
            row["state"].title(),
            f"{row['composite_risk_score']:.1f}",
            f"{row['pn_ratio']:.6f}",
            f"{row['in_ratio']:.3f}",
            f"{row['symptomatic_no_care_pct']:.1f}",
            f"{row['poverty_pct']:.1f}",
        ]
        for _, row in df.sort_values("priority_rank").iterrows()
    ]
    add_table(doc, "Supplementary Table S4. Full state-level prioritization metrics.", columns, rows)


def add_figures(doc: Document) -> None:
    for title, path in FIGURES:
        doc.add_heading(title, level=1)
        if path.exists():
            doc.add_picture(str(path), width=Inches(6))
            caption = doc.add_paragraph(f"{title} Source: {path.name}")
            caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            caption.runs[0].italic = True
        else:
            placeholder = doc.add_paragraph(f"[Missing figure: {path}]")
            placeholder.runs[0].italic = True
        doc.add_page_break()


def create_supplementary_doc() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_layout(doc)
    doc.add_heading("Supplementary Material – Multi-Method TB Detection Delay Analysis", level=0)
    doc.add_paragraph(
        "This file provides extended data tables, PCA outputs, causal pathways, and additional figures referenced in the manuscript."
    )

    component_scores = pd.read_csv(PROCESSED_DIR / "pca_component_scores_delay_determinants.csv")
    loadings = pd.read_csv(PROCESSED_DIR / "pca_loadings_delay_determinants.csv")
    dag_paths = pd.read_csv(PROCESSED_DIR / "dag_causal_paths_delay.csv")
    ranking = pd.read_csv(PROCESSED_DIR / "integrated_state_ranking.csv")

    add_component_scores(doc, component_scores)
    add_loadings(doc, loadings)
    add_dag_paths(doc, dag_paths)
    add_state_rankings(doc, ranking)
    add_figures(doc)

    doc.save(str(OUTPUT_PATH))
    print(f"[SUCCESS] Supplementary material saved to {OUTPUT_PATH}")


if __name__ == "__main__":
    create_supplementary_doc()
