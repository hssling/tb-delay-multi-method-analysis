"""Test image embedding functionality."""
from pathlib import Path
from docx import Document
from docx.shared import Inches

def test_image_embedding():
    """Test embedding images in DOCX."""

    doc = Document()
    doc.add_heading('Image Embedding Test', level=1)

    # Test each image
    images = [
        'forest_plot.png',
        'gis_bubble_map.png',
        'innovative_multipanel_dashboard.png',
        'cluster_map.png'
    ]

    for img_name in images:
        img_path = Path('output/figures') / img_name
        print(f"Testing: {img_path}")

        if img_path.exists():
            try:
                # Add image to doc
                doc.add_paragraph(f"Figure: {img_name}")
                doc.add_picture(str(img_path), width=Inches(5))

                # Add new paragraph with text after image
                doc.add_paragraph(f"Successfully embedded: {img_name}")

                print(f"✓ Successfully embedded: {img_name}")

            except Exception as e:
                print(f"✗ Failed to embed {img_name}: {e}")
                doc.add_paragraph(f"Failed to embed: {img_name} - {e}")
        else:
            print(f"✗ File not found: {img_path}")
            doc.add_paragraph(f"File not found: {img_name}")

    # Save test document
    output_path = Path('reports') / 'test_image_embedding.docx'
    doc.save(output_path)

    print(f"\n✨ Test document saved to: {output_path}")
    print("👀 Check this file to verify images are embedded!")
    print("\nIf images appear, the embedding works!")
    print("If you see only text/error messages, there's an issue.")

if __name__ == "__main__":
    test_image_embedding()
