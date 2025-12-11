#!/usr/bin/env python3
"""
Generate a DOCX containing the key IJMR tables plus a representative figure so the
upload portal has a dedicated tables-and-images document. Existing files remain untouched.
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

PROCESSED_DATA_DIR = Path("data/processed")
OUTPUT_DIR = Path("IJMR_Submission_v3_best")
OUTPUT_PATH = OUTPUT_DIR / "ijmr_tables_and_figures_v2.docx"
FIGURES = [
    (
        "Figure 1. Bayesian posterior distribution of total TB detection delay.",
        "Forest plot generated from the hierarchical NumPyro model showing total delay "
        "posterior means and 95% HDIs across included studies.",
        Path("submission_package/figures/bayesian_forest_total_delay_(days).png"),
    ),
    (
        "Figure 2. Principal component analysis scree plot.",
        "Displays eigenvalues for the eight standardized determinants, highlighting the "
        "three components exceeding the Kaiser criterion.",
        Path("submission_package/figures/pca_scree_plot_delay_determinants.png"),
    ),
    (
        "Figure 3. Directed acyclic graph for TB detection delays.",
        "Summarizes the causal structure linking poverty, care-seeking, and diagnostic "
        "capacity to prevalence-to-notification ratios.",
        Path("submission_package/figures/dag_causal_delay_analysis.png"),
    ),
]


def load_tables():
    bayesian = pd.read_csv(PROCESSED_DATA_DIR / "bayesian_meta_analysis_summary.csv")
    ranking = pd.read_csv(PROCESSED_DATA_DIR / "integrated_state_ranking.csv")
    return bayesian, ranking


def set_layout(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def add_table(doc: Document, caption: str, columns: Sequence[str], rows: Iterable[Sequence[str]]):
    para = doc.add_paragraph(caption)
    para.runs[0].bold = True
    table = doc.add_table(rows=len(rows) + 1, cols=len(columns))
    table.style = "Light List Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, header in enumerate(columns):
        cell = table.cell(0, j)
        cell.text = str(header)
        cell.paragraphs[0].runs[0].bold = True
    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            table.cell(i, j).text = str(value)
    doc.add_paragraph()


def add_bayesian_table(doc: Document, bayesian: pd.DataFrame) -> None:
    columns = ["Delay component", "Mean (days)", "95% HDI (days)", "Studies"]
    rows = [
        [
            row["Delay Type"],
            f"{row['Mean (days)']:.1f}",
            f"{row['HDI 2.5%']:.1f}–{row['HDI 97.5%']:.1f}",
            int(row["Studies"]),
        ]
        for _, row in bayesian.iterrows()
    ]
    add_table(doc, "Table 1. Bayesian posterior estimates for TB detection delays.", columns, rows)


def add_state_table(doc: Document, ranking: pd.DataFrame) -> None:
    top_states = ranking.nsmallest(10, "priority_rank")
    columns = [
        "Priority rank",
        "State",
        "Composite score",
        "Symptomatic no-care (%)",
        "Poverty (%)",
    ]
    rows = [
        [
            int(row["priority_rank"]),
            row["state"].title(),
            f"{row['composite_risk_score']:.1f}",
            f"{row['symptomatic_no_care_pct']:.1f}",
            f"{row['poverty_pct']:.1f}",
        ]
        for _, row in top_states.iterrows()
    ]
    add_table(doc, "Table 2. Top ten states requiring urgent delay mitigation.", columns, rows)


def add_figures(doc: Document) -> None:
    for title, description, image_path in FIGURES:
        doc.add_heading(title, level=1)
        doc.add_paragraph(description)
        if image_path.exists():
            doc.add_picture(str(image_path), width=Inches(6))
            caption = doc.add_paragraph(f"{title} Source: {image_path.name}")
            caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            caption.runs[0].italic = True
        else:
            placeholder = doc.add_paragraph(
                f"[Placeholder – {image_path} not found in repository.]"
            )
            placeholder.runs[0].italic = True
        doc.add_page_break()


def create_tables_and_figures_doc():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_layout(doc)
    doc.add_heading("IJMR Tables and Figures – Tuberculosis Detection Delay Analysis", level=0)
    doc.add_paragraph(
        "Standalone document containing the primary tables and representative figure for IJMR upload."
    )
    bayesian, ranking = load_tables()
    add_bayesian_table(doc, bayesian)
    add_state_table(doc, ranking)
    add_figures(doc)
    doc.save(str(OUTPUT_PATH))
    print(f"[SUCCESS] Tables and figures document saved to {OUTPUT_PATH}")


if __name__ == "__main__":
    create_tables_and_figures_doc()
