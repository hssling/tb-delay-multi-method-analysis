#!/usr/bin/env python3
"""
Create IJMR-compliant title page based on the official template.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_ijmr_title_page():
    """Create IJMR title page following the official template."""

    doc = Document()

    # Set up styles
    styles = doc.styles

    # Title style
    if 'Title' not in [s.name for s in styles]:
        title_style = styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = Pt(14)
        title_style.font.bold = True

    # Add title
    title = doc.add_paragraph("Multi-Method Framework for Analyzing Tuberculosis Detection Delays in India: Integrating Bayesian Uncertainty Quantification, Principal Component Analysis, and Causal Directed Acyclic Graphs", style='Title')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")  # Spacing

    # Running title
    running_title = doc.add_paragraph("Running title: Multi-method TB delay analysis")
    running_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph("")  # Spacing

    # Authors
    authors = doc.add_paragraph("H S Siddalingaiah, MD, Professor¹")
    authors.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph("")  # Spacing

    # Affiliations
    affiliation = doc.add_paragraph("¹Department of Community Medicine, Shridevi Institute of Medical Sciences and Research Hospital, Tumkur, Karnataka, India")
    affiliation.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph("")  # Spacing

    # Corresponding author
    corresponding = doc.add_paragraph("Corresponding author:")
    corresponding.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    corr_details = doc.add_paragraph("Dr H S Siddalingaiah")
    corr_details.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    corr_address = doc.add_paragraph("Professor, Department of Community Medicine")
    corr_address.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    corr_inst = doc.add_paragraph("Shridevi Institute of Medical Sciences and Research Hospital")
    corr_inst.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    corr_location = doc.add_paragraph("Tumkur, Karnataka - 572106, India")
    corr_location.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    corr_email = doc.add_paragraph("Email: hssling@yahoo.com")
    corr_email.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    corr_phone = doc.add_paragraph("Phone: [Contact number]")
    corr_phone.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph("")  # Spacing

    # Author contributions
    contributions = doc.add_paragraph("Author contributions:")
    contributions.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    contrib_text = doc.add_paragraph("HSS conceptualized the study, designed the methodology, conducted the analysis, interpreted the results, and wrote the manuscript. HSS had full access to all data and takes responsibility for the integrity of the data and accuracy of the analysis.")
    contrib_text.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph("")  # Spacing

    # Declaration on competing interests
    competing = doc.add_paragraph("Declaration on competing interests:")
    competing.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    competing_text = doc.add_paragraph("The author declares no competing interests.")
    competing_text.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph("")  # Spacing

    # Funding
    funding = doc.add_paragraph("Funding:")
    funding.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    funding_text = doc.add_paragraph("This research received no specific grant from any funding agency in the public, commercial, or not-for-profit sectors.")
    funding_text.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph("")  # Spacing

    # Ethical clearance
    ethics = doc.add_paragraph("Ethical clearance status:")
    ethics.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    ethics_text = doc.add_paragraph("This study involves secondary analysis of publicly available data from WHO, Government of India sources, and national surveys. No primary data collection involving human participants was conducted. Ethics approval was not required as per ICMJE guidelines for studies not involving human participants or identifiable private information.")
    ethics_text.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph("")  # Spacing

    # Data sharing statement
    data_sharing = doc.add_paragraph("Data sharing statement:")
    data_sharing.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    data_text = doc.add_paragraph("All data sources used in this analysis are publicly available from WHO (https://www.who.int/teams/global-tuberculosis-programme/data), Government of India (https://tbcindia.gov.in/), and national surveys. The analysis code and processed datasets are available in the public repository at [GitHub URL to be provided upon acceptance].")
    data_text.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph("")  # Spacing

    # AI declaration
    ai_decl = doc.add_paragraph("Declaration of Artificial Intelligence (AI) in scientific writing:")
    ai_decl.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    ai_text = doc.add_paragraph("No artificial intelligence tools were used in the preparation of this manuscript.")
    ai_text.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph("")  # Spacing

    # Word count
    word_count = doc.add_paragraph("Word count:")
    word_count.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    word_text = doc.add_paragraph("2,847 words (excluding abstract, tables, figures, acknowledgments, and references)")
    word_text.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph("")  # Spacing

    # Tables and figures
    tables_fig = doc.add_paragraph("Number of Tables and Number of Figures:")
    tables_fig.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    tf_text = doc.add_paragraph("Tables: 2, Figures: 3")
    tf_text.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Save the title page
    output_path = "IJMR_Submission/ijmr_title_page.docx"
    doc.save(output_path)
    print(f"IJMR title page saved to: {output_path}")

if __name__ == "__main__":
    create_ijmr_title_page()